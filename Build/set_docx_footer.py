#!/usr/bin/env python3
"""Setzt eine einheitliche Fußzeile in einem von Pandoc erzeugten DOCX."""

from pathlib import Path
import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--text", required=True)
    args = parser.parse_args()

    document = Document(args.docx)

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(args.text)
        run.font.size = Pt(8)

    document.save(args.docx)


if __name__ == "__main__":
    main()
